# services/file_parser.py
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
import pytesseract
from PIL import Image
import logging

logger = logging.getLogger("fileparser")


class FileParser:
    """
    Production-grade parser.
    - parse()        → structured text for translation
    - parse_text_only() → safe text-only extraction (NO OCR)
    - parse_text_with_ocr() → explicit OCR path (glossary only)
    """

    # =========================
    # NORMAL TRANSLATION PATH
    # =========================
    def parse(self, filepath: Path):
        filepath = Path(filepath)
        suffix = filepath.suffix.lower()

        if suffix == ".pdf":
            return self._parse_pdf(filepath)
        if suffix == ".docx":
            return self._parse_docx(filepath)
        if suffix == ".txt":
            return self._parse_txt(filepath)

        return {"text": "", "pages": 1, "positions": None}

    def _parse_txt(self, path: Path):
        text = path.read_text(encoding="utf-8", errors="ignore")
        return {"text": text, "pages": 1, "positions": None}

    def _parse_docx(self, path: Path):
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return {"text": text, "pages": 1, "positions": None}

    def _parse_pdf(self, path: Path):
        doc = fitz.open(str(path))
        all_texts = []
        positions = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("blocks")

            for b in blocks:
                x0, y0, x1, y1, text, *_ = b
                t = text.strip()
                if not t:
                    continue

                all_texts.append(t)
                positions.append({
                    "text": t,
                    "bbox": [x0, y0, x1 - x0, y1 - y0],
                    "page": page_num + 1
                })

        return {
            "text_blocks": all_texts,
            "positions": positions,
            "pages": len(doc)
        }

    # =========================
    # SAFE TEXT-ONLY (NO OCR)
    # =========================
    def parse_text_only(self, filepath: Path) -> str:
        suffix = filepath.suffix.lower()

        if suffix == ".txt":
            return filepath.read_text(errors="ignore")

        if suffix == ".docx":
            doc = Document(str(filepath))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        if suffix == ".pdf":
            try:
                doc = fitz.open(str(filepath))
                text = ""
                for page in doc:
                    text += page.get_text()
                return text
            except Exception as e:
                logger.warning(f"PDF text extraction failed: {e}")
                return ""

        return ""

    # =========================
    # OCR — EXPLICIT ONLY
    # =========================
    def parse_text_with_ocr(self, filepath: Path) -> str:
        suffix = filepath.suffix.lower()

        # PDF OCR
        if suffix == ".pdf":
            try:
                from pdf2image import convert_from_path
                images = convert_from_path(str(filepath))
                text = ""
                for img in images:
                    text += pytesseract.image_to_string(img, lang="rus+eng")
                return text
            except Exception as e:
                logger.error(f"OCR failed: {e}")
                return ""

        # Image OCR
        try:
            img = Image.open(str(filepath)).convert("RGB")
            return pytesseract.image_to_string(img, lang="rus+eng")
        except:
            return ""
